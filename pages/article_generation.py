from dotenv import load_dotenv
import os
from langgraph.graph import StateGraph, START, END
from typing import TypedDict, Literal, Annotated
from langchain_openai import ChatOpenAI
import streamlit as st
from pydantic import BaseModel, Field
import operator
import time
from docx import Document
import io

load_dotenv(dotenv_path=r"my_api_key.env")
os.environ['OPENAI_API_KEY'] = os.getenv("openai_api")

os.environ["LANGCHAIN_TRACING_V2"] = 'true'
os.environ["LANGCHAIN_ENDPOINT"] = "https://api.smith.langchain.com"
os.environ["LANGCHAIN_API_KEY"] = os.getenv("langsmith_api")
os.environ["LANGCHAIN_PROJECT"] = 'article_generation'

if 'login' not in st.session_state:
    st.session_state['login'] = False


elif st.session_state['login'] == True:

    # loading model for different purpose
    model = ChatOpenAI(model="gpt-4.1-mini")
    evaluate_llm = ChatOpenAI(model="gpt-3.5-turbo")
    optimize_llm = ChatOpenAI(model='gpt-5-nano-2025-08-07')

    # creating a model structure

    class evaluationmodel(BaseModel):
        evaluation: Literal['approved', 'not approved'] = Field(
            decription="evaluation of the article")
        feedback: str = Field(description="feedback for the article")

    # giving the evaluate model a structure
    structured_evaluate = evaluate_llm.with_structured_output(evaluationmodel)

    # creating a state model

    class textstate(TypedDict):
        question: str
        content: str
        keyword: str
        evaluate: Literal['approved', 'not_approved']
        feedback: str
        iteration: int
        max_iteration: int
        feedback_hsitory: Annotated[list[str], operator.add]
        article_hist: Annotated[list[str], operator.add]

    # creating a artile generation graph

    def create_content(state: textstate):
        question = state['question']

        prompt = f"""write an article on the given topic {question} in a very professional manner .
            The article written must include the important details from the topic and must be liked by everyone who reads it . 
            Make Sure that you do not include any unnecessary things in the article.
            You have to do a proper research work and gather each and every important details for the topic.
            make sure that the content is knowledgeable and  propvides a useful information.
            Also specify the target audience that who might be interested in reading the article on the topic specified. 
            ensure to write in a article format only.
            Make sure to replace the normal words with seo keywords whenever required.
            """
        response = model.invoke(prompt).content

        return {'content': response, 'article_hist': [response]}

    def evaluate_content(state: textstate):
        prompt = f"""
                you have to evaluate the content generated on topic {state['question']} and 
                the content {state['content']}.
                Based on the above topic and content you have to evaluate that wheather the article given must be approved or  not . 
                based on the criteria you have to evaluate the article. 
                if it is written in question - answer format it must be rejected , 
                must highlight the keywords and important note.
                make sure to create an summary line at the end of the article.
                on the basis of this give the :
                evaluation :"approved" or "not_approved"
                feedback : in one line explaning strength and weakness 
                """
        response = structured_evaluate.invoke(prompt).content

        return {'evaluate': response.evaluation, 'feedback': response.feedback, 'feedback_history': response.feedback}

    def optimizing_content(state: textstate):
        prompt = f"""
            improve the article on the basis of the feedback 
            {state['content']} and {state['feedback']}
            you have to re-write the article if they again ask for the same topic to generate.
        """
        response = optimize_llm.invoke(prompt)
        iteration = state['iteration']+1

        return {'content': response, 'iteration': iteration, 'article_hist': [response]}

    # routing the evaluation

    def route_evaluation(state: textstate):
        if state['evaluate'] == 'approved' or state['iteration'] >= state['max_iteration']:
            return 'approved'
        else:
            return 'not_approved'

    # creating graph and compiling
    # used a looping pattern to evaluate the content and provide with the better results
    graph = StateGraph(textstate)

    # creating nodes
    graph.add_node("create_content", create_content)
    graph.add_node('evaluate', evaluate_content)
    graph.add_node('optimize', optimizing_content)

    # creating edges
    graph.add_edge(START, "create_content")
    graph.add_conditional_edges('evaluate', route_evaluation, {
                                'approved': END, 'not_approved': 'optimize'})
    graph.add_edge('optimize', END)

    # compiling
    workflow = graph.compile()

    # STREAMLIT
    st.header('ARTICLE GENERATION')
    st.markdown('----------')
    # initial state
    # taking input
    input = st.text_input(
        label="GENERATE", placeholder="enter the topic for the article")

    # initiating the input and invoking it
    initial_state = {
        "question": input,
        'iteration': 1,
        'max_iteration': 3
    }

    # creating a button to generate
    generate = st.button('GENERATE ARTICLE', use_container_width=True)

    if generate:
        with st.spinner('GENERATING RESPONSE....'):
            time.sleep(5)
            final_state = workflow.invoke(initial_state)
            article = final_state['content']

        st.session_state['article'] = article
        st.markdown(st.session_state['article'])

    # adding the generated article in docx file
    insert_into_file = st.button("INSERT INTO FILE", use_container_width=True)
    file_name = initial_state['question']

    if insert_into_file:
        if 'article' not in st.session_state:
            st.error("please generate the content first")
        else:
            try:
                doc = Document(f"{file_name}.docx")

            except Exception:
                doc = Document()

            doc.add_paragraph(st.session_state['article'])
            bio = io.BytesIO()
            # saving the doc
            doc.save(bio)
            bio.seek(0)
            st.success("data inserted successfully!!")

            # creating a downloa button to download the file
            st.download_button(
                label="DOWNLOAD FILE",
                data=bio.getvalue(),
                file_name=f'{file_name}_article.docx',
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                icon=":material/download:",
                width='stretch'
            )

    html = st.button("GENERATE HTML", use_container_width=True)
    if html:
        st.switch_page("pages/article_html.py")

else:
    st.switch_page("password.py")
